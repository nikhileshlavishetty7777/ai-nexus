# backend/services/file_service.py
import os
import uuid
import aiofiles
from typing import Optional
from fastapi import UploadFile, HTTPException
from sqlalchemy.orm import Session
from database.models import Document
from config import settings
from loguru import logger


ALLOWED_TYPES = {
    "application/pdf": "pdf",
    "text/plain": "txt",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword": "doc",
}


async def save_upload_file(file: UploadFile, user_id: str) -> tuple[str, str]:
    """Save uploaded file and return (filepath, file_type)."""
    content_type = file.content_type or ""
    file_type = ALLOWED_TYPES.get(content_type)

    if not file_type:
        # Try by extension
        ext = os.path.splitext(file.filename or "")[-1].lower().lstrip(".")
        if ext in ["pdf", "txt", "docx"]:
            file_type = ext
        else:
            raise HTTPException(status_code=400, detail=f"File type not supported. Allowed: PDF, TXT, DOCX")

    # Check file size
    content = await file.read()
    if len(content) > settings.MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail=f"File too large. Max size: 10MB")

    # Save file
    user_dir = os.path.join(settings.UPLOAD_DIR, user_id)
    os.makedirs(user_dir, exist_ok=True)

    filename = f"{uuid.uuid4()}.{file_type}"
    filepath = os.path.join(user_dir, filename)

    async with aiofiles.open(filepath, "wb") as f:
        await f.write(content)

    return filepath, file_type


async def extract_text(filepath: str, file_type: str) -> str:
    """Extract text content from uploaded file."""
    try:
        if file_type == "txt":
            async with aiofiles.open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                return await f.read()

        elif file_type == "pdf":
            import PyPDF2
            text = []
            with open(filepath, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text.append(page.extract_text() or "")
            return "\n".join(text)

        elif file_type in ["docx", "doc"]:
            import docx
            doc = docx.Document(filepath)
            return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

    except Exception as e:
        logger.error(f"Text extraction error: {e}")
        return ""


async def process_document(db: Session, doc_id: str):
    """Process document: extract text and generate summary."""
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        return

    try:
        text = await extract_text(doc.file_path, doc.file_type)
        doc.extracted_text = text[:50000]  # Limit stored text

        # Generate a brief summary using AI if available
        if text and len(text) > 100:
            try:
                from services.ai_service import quick_ai_action
                summary = await quick_ai_action(
                    "summarize",
                    text[:3000],
                    model="gpt-4o-mini",
                    provider="openai",
                )
                doc.summary = summary[:1000]
            except Exception as e:
                logger.warning(f"Could not generate summary: {e}")

        doc.is_processed = True
        db.commit()
        logger.info(f"Document processed: {doc.original_filename}")
    except Exception as e:
        logger.error(f"Document processing error: {e}")


def create_document_record(
    db: Session,
    user_id: str,
    original_filename: str,
    filename: str,
    file_type: str,
    file_size: int,
    file_path: str,
    chat_id: Optional[str] = None,
) -> Document:
    doc = Document(
        user_id=user_id,
        chat_id=chat_id,
        filename=filename,
        original_filename=original_filename,
        file_type=file_type,
        file_size=file_size,
        file_path=file_path,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    return doc
